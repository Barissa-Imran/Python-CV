# Built by Imran Dhidha Barissa
# https://linkedin.com/in/imran-barissa
# This CV builder is meant to build a cv/resume for a student or recent graduate in that format
# More modifications will be done later to include columns and sections

from docx import Document
from docx.shared import Inches
import pyttsx3 

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'MEE 2.png', 
    width=Inches(2.0)
)

# name phone number and email details
speak('what is your name? ')
name = input('what is your name? ')
speak('Hello' + name +'How are you today? ')

speak('what is your phone number? ')
phone_number = input('what is your phone number? ')

speak('please provide an email address')
email = input('what is your email? ')

document.add_paragraph(
    name + ' | '+ phone_number + ' | ' + email
)
# about me
document.add_heading('About me')
speak('Tell me about yourself?')
document.add_paragraph(
    input('Tell me about yourself? ')
)
# Education background
document.add_heading('Education')
p = document.add_paragraph()

speak('This is the Education section')
speak('Please read the prompts on the screen and answer appropriately')
university = input('Name of the institution '+ '\n')
degree = input('Degree undertook '+'\n')
from_date = input('From date? '+'\n')
to_date = input('To Date? '+'\n')
experience = input('What was your experience? '+'\n')

p.add_run(university + ', ' + degree +' | ').bold = True
p.add_run(from_date + '-' + to_date +'\n').italic = True
p.add_run(experience+'\n')

#More institutions
while True:
    speak('Do you have any more educational background?')
    has_more_education = input('Do you more educational backround? Yes or No')

    if has_more_education.lower() == 'yes':
        P = document.add_paragraph()

        speak('Please add Education as indicated')
        university = input('Name of the institution '+ '\n')
        degree = input('Degree undertook '+'\n')
        from_date = input('From date? '+'\n')
        to_date = input('To Date? '+'\n')
        experience = input('What was your experience? '+'\n')

        p.add_run(university + ', ' + degree +' | ').bold = True
        p.add_run(from_date + '-' + to_date +'\n').italic = True
        p.add_run(experience+'\n')
    else:
        break

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('This is the work experience section')
speak('Please read the prompts on the screen and answer appropiately')
job_title = input('Job title?'+ '\n')
company = input('Enter Company' + '\n' )
from_date = input('From date'+ '\n')
to_date = input('To Date'+ '\n')

p.add_run(job_title+ ', '+ company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic

experience_details = input(
    'Describe your experience at' + ' '+ company + '\n'
)
p.add_run(experience_details)

# more experiences
while True:
    speak ('Do you have more experiences?')
    has_more_experiences = input('Do you have more experiences? Yes or No ')

    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        speak('Please provide your experience following the prompts')
        job_title = input('Job title?'+ '\n')
        company = input('Enter Company '+'\n')
        from_date = input('From date '+'\n')
        to_date = input('To Date'+'\n')

        p.add_run(job_title+ ', '+ company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic

        experience_details = input(
            'Describe your experience at' + ' '+ company +'\n'
        )
        p.add_run(experience_details)
    else:
        break

# Projects
document.add_heading('Projects')
p = document.add_paragraph()

speak('This is the projects section')
speak('Fill in below following the prompts')
project_name = input ('Name of project? '+'\n')
language = input ('Programming language used' + '\n')
repo = input('Github repository link '+ '\n')
description = input('Provide project description '+'\n')

p.add_run(project_name +' | ').bold = True
p.add_run(repo + '\n').italic = True
p.add_run(description + '\n')
p.add_run(language).bold = True

# More projects
while True:
    speak('Do you have more projects')
    has_more_projects = input('Do you have more projects? Yes or No')

    if has_more_projects.lower() == 'yes':
        p = document.add_paragraph()

        project_name = input ('Name of project? '+'\n')
        language = input ('Programming language used' + '\n')
        repo = input('Github repository link '+ '\n')
        description = input('Provide project description '+'\n')

        p.add_run(project_name +' | ').bold = True
        p.add_run(repo + '\n').italic = True
        p.add_run(description + '\n')
        p.add_run(language).bold = True
    else:
        break

# Leadership and Awards
document.add_heading('Leadership and Awards')
p = document.add_paragraph()

speak('This is the leadership and awards section')
speak('fill in appropriately')
award_title = input ('Name of award or leadership position'+ '\n')
institution = input('Providing institution'+'\n')
description = input('describe the purpose of the award')

p.add_run(award_title + ' - ' + institution +'\n').bold = True
p.add_run(description)

#more Leadership and Awards
while True:
    speak('Do you have more')
    has_more_awards = input('Do you have more leadership roles and awards? Yes or No')

    if has_more_awards.lower() == 'yes':
        p = document.add_paragraph()

        award_title = input ('Name of award or leadership position'+ '\n')
        institution = input('Providing institution'+'\n')
        description = input('describe the purpose of the award')

        p.add_run(award_title + ' - ' + institution +'\n').bold = True
        p.add_run(description)
    else:
        break

speak('Proceed to viewing the document you have created')
document.save('cv.docx')
